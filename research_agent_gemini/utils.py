import re
import io
from fpdf import FPDF
from docx import Document

def clean_text(text):
    """
    Removes unsupported characters for PDF generation.
    FPDF2 default fonts don't support emojis, so we strip them or replace them.
    This is a basic cleaner.
    """
    # Remove emojis or non-latin characters that might cause issues with standard fonts
    # For simplicity, we'll strip non-ascii but keep basic latin-1
    # text = text.encode('latin-1', 'ignore').decode('latin-1')
    # Actually fpdf2 supports unicode if we use a ttf font, but we don't have one bundled easily.
    # So we will try to replace common emojis with text or just strip them.

    # Simple strategy: remove characters outside basic ranges if they cause trouble,
    # but let's try to keep as much as possible and handle errors if they arise.
    # For now, let's just return the text and let FPDF2 handle it,
    # but we might need to add a font if we want emoji support.
    # Since we can't easily add a font file without downloading it, we'll strip emojis.

    # Regex to remove emojis (broad range)
    return text.encode('latin-1', 'replace').decode('latin-1')

def generate_pdf(text: str) -> bytes:
    """
    Generates a PDF from the given text with basic Markdown formatting.
    """
    class PDF(FPDF):
        def header(self):
            self.set_font('Helvetica', 'B', 12)
            self.cell(0, 10, 'Research Report', 0, 1, 'C')
            self.ln(5)

        def footer(self):
            self.set_y(-15)
            self.set_font('Helvetica', 'I', 8)
            self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=12)

    # Clean text for PDF compatibility (standard fonts don't support many unicode chars)
    # We replace unsupported chars with ?

    lines = text.split('\n')

    for line in lines:
        line = clean_text(line)
        if not line.strip():
            pdf.ln(5)
            continue

        if line.startswith('# '):
            pdf.set_font("Helvetica", 'B', 16)
            pdf.multi_cell(0, 10, line[2:].strip(), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=12)
        elif line.startswith('## '):
            pdf.set_font("Helvetica", 'B', 14)
            pdf.multi_cell(0, 10, line[3:].strip(), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=12)
        elif line.startswith('### '):
            pdf.set_font("Helvetica", 'B', 13)
            pdf.multi_cell(0, 10, line[4:].strip(), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=12)
        elif line.startswith('* ') or line.startswith('- '):
            pdf.set_x(20) # Indent
            pdf.multi_cell(0, 8, f"\x95 {line[2:].strip()}", new_x="LMARGIN", new_y="NEXT")
        else:
            # Handle bold (simple check)
            # This is a very basic parser, it doesn't handle inline bolding perfectly with multi_cell
            # so we just print plain text for now to avoid complexity.
            # If we wanted bold, we'd need to parse chunks.
            # For this MVP, we strip markdown bold syntax.
            clean_line = line.replace('**', '').replace('__', '')
            pdf.multi_cell(0, 8, clean_line, new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())

def generate_docx(text: str) -> io.BytesIO:
    """
    Generates a DOCX from the given text with basic Markdown formatting.
    """
    doc = Document()
    doc.add_heading('Research Report', 0)

    lines = text.split('\n')

    for line in lines:
        if not line.strip():
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
        else:
            # Handle bold (simple check)
            p = doc.add_paragraph()
            # Simple bold parser: split by **
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
